###################################
# !/usr/bin/python
# Python 3.10
# (C) 2023 admi.tech
###################################

# Import main packages
import sys
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import csv


def create_file(myData):
    """This function will take dictionary and create full document."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Word Header

    section = doc.sections[0]
    header = section.header
    par = header.paragraphs[0]
    par_format = par.paragraph_format
    par_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.text = f"{myData['firma']}, Hviezdoslavova 315/35, 905 01 Senica\n ----------------------------------------------------------------------"

    # Heading
    heading1 = doc.add_heading('DOCHÁDZKA')
    heading_format = heading1.paragraph_format
    heading_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Heading
    heading2 = doc.add_heading('Evidencia pracovného času zamestnanca', level=2)
    heading_format = heading2.paragraph_format
    heading_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #
    par1 = doc.add_paragraph()
    paragraph_format = par1.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par1.add_run(f'\n Meno a Priezvisko : {myData["meno"]} {myData["priezvisko"]}')

    """
    # Miesto a datum
    par1 = doc.add_paragraph()
    paragraph_format = par1.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    par1.add_run('\n')
    par1.add_run('\t\t\t\t\t\t\t\tV Senici '+str(now.day)+'.'+str(now.month)+'.'+str(now.year))


    # Adresa Cloviecika
    if pohlavie(Meno) == 'F':
        Oslovenie = "Vážená pani"
    else:
        Oslovenie = "Vážený pán"
    par2 = doc.add_paragraph('\n\n\t\t\t\t\t\t\t\t'+Oslovenie+'\n')
    par2.add_run('\t\t\t\t\t\t\t\t'+Meno+' '+Priezvisko+'\n')
    par2.add_run('\t\t\t\t\t\t\t\t'+Ulica+'\n')
    par2.add_run('\t\t\t\t\t\t\t\t'+str(Psc)+' '+Mesto+'\n')
    par2.add_run('\t\t\t\t\t\t\t\t'+Stat+'\n')
    par2.add_run('\n')

    # par3 = doc.add_paragraph('text '+Mesto).underline
    # par3.add_run('\n\n')
    par3 = doc.add_paragraph()
    run = par3.add_run()
    run.font.bold = True
    run.font.underline = True
    run.text = 'Vtext '

    """

    # Document Save

    doc.save(f"{myData['meno']}_{myData['priezvisko']}_{myData['mesiac']}_{myData['rok']}.docx")



# Run the example

if __name__ == "__main__":

    sample_data = {
        "meno": "Peter",
        "priezvisko": "Smith",
        "firma": "Microsoft",
        "mesiac": 5,
        "rok": 2023,
        "dochadzka": {
            '2.5.': [1, '7:23', '11:33', '12:03', '16:00'],
            '3.5.': [2, '7:32', '11:26', '11:56', '16:02'],
            '4.5.': [3, '7:25', '11:32', '12:02', '16:00'],
            '5.5.': [4, '7:20', '11:16', '11:46', '16:00'],
            '6.5.': [5, '7:28', '11:32', '12:02', '16:00'],
            '7.5.': [6, '7:31', '11:43', '12:13', '16:01'],
            '8.5.': 'Sviatok',
        }
    }

    create_file(sample_data)
